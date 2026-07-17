from __future__ import annotations
from io import BytesIO
from pathlib import Path
from pypdf import PdfReader
from docx import Document

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}

class ExtractionError(RuntimeError):
    pass

def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED:
        raise ExtractionError(f"지원하지 않는 파일 형식입니다: {suffix}")

    try:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(data))
            pages = []
            for index, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                pages.append(f"\\n[페이지 {index}]\\n{text}")
            result = "\\n".join(pages)
        elif suffix == ".docx":
            doc = Document(BytesIO(data))
            result = "\\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for table_index, table in enumerate(doc.tables, start=1):
                result += f"\\n[표 {table_index}]\\n"
                for row in table.rows:
                    result += " | ".join(cell.text.strip() for cell in row.cells) + "\\n"
        else:
            result = data.decode("utf-8-sig", errors="replace")
    except Exception as exc:
        raise ExtractionError(f"문서 추출 실패: {exc}") from exc

    if not result.strip():
        raise ExtractionError("문서에서 텍스트를 추출하지 못했습니다. 스캔 PDF는 OCR이 필요합니다.")
    return result.strip()
